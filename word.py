from docx import Document

doc = Document()

sections = {
"Behavioral Questions": {
"Tell me about yourself.": 
"I am a Master's student in CSE at the University of Michigan with experience in embedded software development at Mercedes-Benz R&D and AI/ML research under Prof. Atul Prakash. I enjoy hands-on work, solving practical problems, and supporting people with technology. I bring strong troubleshooting skills, clear communication, and a calm, organized working style.",
"Why are you interested in working at SLTS?":
"SLTS provides hands-on troubleshooting experience across 75+ campus buildings. I enjoy supporting AV systems, digital signage, tablets, and access devices while interacting with staff and students. I like the balance of independent field work and teamwork, and I want to contribute to the reliability of Student Life infrastructure.",
"What qualities or skills do you have that relate to this position?":
"I bring calm analytical thinking, systematic troubleshooting, strong communication, and patience. I am reliable, punctual, and organized. My engineering experience taught me to diagnose issues carefully, document thoroughly, and stay consistent.",
"What sets you apart from other candidates?":
"I think in systems rather than individual issues, stay calm under pressure, and adapt quickly to unfamiliar equipment. My mix of industry experience, research work, and hands-on troubleshooting helps me approach problems logically and communicate clearly."
},

"Communication & Customer Service": {
"Describe a time you explained a complex issue to a non-technical person.":
"At Mercedes-Benz, I explained inconsistent software test results to a business manager by focusing on impact rather than technical details. I said the verification tool was giving unreliable information and moving forward could cause bigger issues. I provided a clear plan and timeline, which kept communication simple and professional.",
"How would you handle an upset or stressed user?":
"I stay calm, listen carefully, acknowledge their frustration, and reassure them I am there to help. I avoid jargon, communicate clearly, and provide alternatives when possible. My goal is to help them feel supported and informed.",
"How do you explain technical issues without jargon?":
"I simplify by focusing on impact, such as: 'This device is not sending a signal, so the display cannot show anything.' I confirm understanding before continuing."
},

"Scenario-Based Troubleshooting": {
"A projector is not displaying. What do you do?":
"I check power, cable connections, and ensure the correct input source. I test with another device or cable and restart the projector. If the issue continues, I document all steps and escalate if needed.",
"A digital sign is frozen or blank.":
"I confirm power, check HDMI and content player status, inspect network connectivity, and perform a power cycle. I verify content scheduling if needed.",
"A staff laptop won’t connect to Wi-Fi.":
"I ensure WiFi is enabled, reconnect or forget the network, restart the adapter, test other devices, and check for local outages.",
"A card reader at a building entrance isn’t responding.":
"I check indicator lights, confirm power, test multiple cards, record the serial number, document steps, and escalate to the access control team."
},

"Teamwork & Independence": {
"Describe a time you worked independently.":
"During my internship at Mercedes-Benz, I handled long debugging cycles independently. I maintained logs, communicated progress, and stayed disciplined. This strengthened my ability to work without supervision.",
"Describe a time you worked in a team.":
"In my CVD AI project, my team divided tasks and cross-verified each other's work. Our collaboration allowed deployment across five hospitals."
},

"Work Habits & Fit": {
"How do you stay organized?":
"I use time‑blocking, checklists, and calendars. I break tasks into small steps and adjust priorities as new tasks come in.",
"How do you handle two tasks arriving at the same time?":
"I evaluate urgency, communicate timelines, address the critical task first, document progress, and follow up quickly on the second task.",
"Why should we hire you?":
"I combine strong technical troubleshooting skills with calm communication and reliability. I enjoy helping people, learning new systems, and supporting campus infrastructure. I can work independently, document well, and stay consistent."
},

"Technical Problem Solving": {
"Describe a technical problem you solved that wasn’t obvious.":
"During my AURA research, the agent entered repeated reasoning loops with no errors. I added detailed tracing, isolated each step, and discovered an indexing issue in the memory module. I redesigned the logic and added integrity checks, which resolved the problem.",
"Describe a mistake you made and how you fixed it.":
"I once misinterpreted a documentation requirement. I informed my supervisor, corrected it, and adopted a habit of reconfirming instructions to avoid future mistakes."
},

"Quick Summary Page (Cheat Sheet)": {
"Key Strengths":
"Calm • Systematic • Clear communication • Reliable • Independent worker • Patient • Fast learner.",
"Fast Troubleshooting Approach":
"Check power • Check cables • Restart • Test alternate device • Document • Escalate logically.",
"Customer Service Approach":
"Stay calm • Listen • Acknowledge frustration • Avoid jargon • Offer alternatives • Communicate clearly."
}
}

for section, qa in sections.items():
    doc.add_heading(section, level=1)
    for q, a in qa.items():
        doc.add_heading(q, level=2)
        doc.add_paragraph(a)

filepath = "SLTS_Updated_Important_Interview_Prep_Bhargav.docx"
doc.save(filepath)

filepath
